from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # --- Header / Title ---
    title = doc.add_heading('Cyclistic: Behavioral Mirroring Model', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Strategic Member Conversion via Behavioral DNA')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph() # Spacer

    # --- Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    p = doc.add_paragraph(
        "This report outlines the development and findings of the Behavioral Mirroring Model, "
        "a data-driven framework designed to optimize Cyclistic's member conversion strategy. "
        "Traditional marketing efforts often rely on high-traffic volume, which can be inefficient. "
        "Our model shifts the focus to 'velocity of habit'—identifying stations where casual riders "
        "already exhibit the 'Twin Peaks' commuting patterns characteristic of annual members."
    )

    # --- Business Background ---
    doc.add_heading('2. Business Background & Goal', level=1)
    doc.add_paragraph(
        "Cyclistic, a bike-share program in Chicago, seeks to maximize the number of annual memberships. "
        "Annual members are more profitable than casual riders. The core business question addressed is:"
    )
    doc.add_paragraph(
        "\"How do annual members and casual riders use Cyclistic bikes differently, and how can we design "
        "a marketing strategy to convert casual riders into annual members?\"",
        style='Quote'
    )

    # --- Methodology ---
    doc.add_heading('3. Methodology: The Behavioral Mirroring Model', level=1)
    doc.add_paragraph(
        "The model analyzes hourly ridership 'DNA' for every station in the network. Key indicators include:"
    )
    doc.add_paragraph("Twin Peaks Signal: Identifying rush-hour spikes at 8 AM and 5 PM.", style='List Bullet')
    doc.add_paragraph("Behavioral Mirroring: Comparing casual rider behavior against established member patterns.", style='List Bullet')
    doc.add_paragraph("Pearson Correlation: Statistically validating behavioral similarity at potential 'Anchor' stations.", style='List Bullet')

    # --- Technical Pipeline ---
    doc.add_heading('4. Technical Pipeline Overview', level=1)
    doc.add_paragraph(
        "The analysis follows a modular 6-stage Jupyter Notebook pipeline ensuring data integrity and deep insight:"
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Stage'
    hdr_cells[1].text = 'Purpose'
    
    stages = [
        ("1. Data Pipeline", "Ingestion, cleaning, and schema normalization."),
        ("2. Habitual Analysis", "Calculation of station-level 'Routine Scores'."),
        ("3. Behavioral Refinement", "Classification into mirror-verdict categories."),
        ("4. Station Segmentation", "Final portfolio categorization (Anchors, Emerging, Noise)."),
        ("5. Mirror Correlation", "Pearson correlation validation of high-value targets."),
        ("6. Visual Generation", "Production-ready charts and stakeholder visuals.")
    ]
    
    for stage, purpose in stages:
        row_cells = table.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = purpose

    # --- Key Insights ---
    doc.add_heading('5. Key Insights & Results', level=1)
    doc.add_paragraph(
        "The analysis segmented over 1,500 bike stations into three actionable categories:"
    )
    doc.add_paragraph("Confirmed Anchors: Primary conversion targets with 67% member share during peaks.", style='List Bullet')
    doc.add_paragraph("High-Potential Emerging: Stations showing intermittent mirroring signals.", style='List Bullet')
    doc.add_paragraph("Inconsistent / Noise: Tourist-heavy stations with low commuting signals.", style='List Bullet')
    
    doc.add_paragraph(
        "The 'Twin Peaks Signal' was validated across 13 months of trip data (Dec 2024 – Dec 2025)."
    )

    # --- Strategic Recommendations ---
    doc.add_heading('6. Strategic Recommendations', level=1)
    doc.add_paragraph(
        "To maximize ROI, Cyclistic should focus marketing efforts on Behavioral Anchor stations. "
        "Recommendations include:"
    )
    doc.add_paragraph("Targeted On-Site Activation: Launch campaigns at Anchor stations during 7–9 AM and 4–6 PM windows.", style='List Bullet')
    doc.add_paragraph("Digital Precision: Use station-level data to trigger localized digital ads for casual riders in Anchor clusters.", style='List Bullet')
    doc.add_paragraph("Dynamic Pricing: Trial 'Member-Beta' passes at Anchor stations to bridge the gap between casual and member behavior.", style='List Bullet')

    # --- Conclusion ---
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        "By moving from volume-based targeting to behavioral mirroring, Cyclistic can achieve higher conversion "
        "rates with lower marketing spend. The Behavioral Mirroring Model provides the technical foundation for "
        "this data-driven shift in strategy."
    )

    # Save the document
    doc_path = "Cyclistic_Project_Report.docx"
    doc.save(doc_path)
    print(f"Report successfully saved to {doc_path}")

if __name__ == "__main__":
    create_report()
